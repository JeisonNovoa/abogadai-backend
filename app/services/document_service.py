from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def _agregar_numero_pagina(canvas, doc):
    """
    Callback para agregar número de página al pie
    """
    canvas.saveState()
    page_num = canvas.getPageNumber()
    text = f"Página {page_num}"
    canvas.setFont('Times-Roman', 9)
    canvas.drawCentredString(4.25 * inch, 0.5 * inch, text)
    canvas.restoreState()


def generar_pdf(documento_texto: str, nombre_solicitante: str = "Documento") -> BytesIO:
    """
    Genera un PDF del documento legal con formato profesional
    """
    buffer = BytesIO()

    # Crear el documento PDF
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=54,  # Espacio para numeración
    )

    # Estilos
    styles = getSampleStyleSheet()

    # Estilo para título principal (ACCIÓN DE TUTELA, DERECHO DE PETICIÓN)
    styles.add(ParagraphStyle(
        name='TituloPrincipal',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=14,
        leading=18,
        spaceAfter=12,
        spaceBefore=0,
        fontName='Times-Bold',
    ))

    # Estilo para títulos de secciones (I., II., III., etc.)
    styles.add(ParagraphStyle(
        name='TituloSeccion',
        parent=styles['Heading2'],
        alignment=TA_LEFT,
        fontSize=12,
        leading=16,
        spaceAfter=8,
        spaceBefore=12,
        fontName='Times-Bold',
    ))

    # Estilo para texto justificado
    styles.add(ParagraphStyle(
        name='Justify',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14,
        fontName='Times-Roman',
        spaceAfter=6,
    ))

    # Estilo para encabezado formal (Señor, JUEZ, etc.)
    styles.add(ParagraphStyle(
        name='Encabezado',
        parent=styles['BodyText'],
        alignment=TA_LEFT,
        fontSize=11,
        leading=14,
        fontName='Times-Roman',
        spaceAfter=4,
    ))

    # Contenido
    story = []

    # Dividir el texto en párrafos
    lineas = documento_texto.split('\n')

    # Patrones para detectar diferentes tipos de formato
    patron_titulo_principal = r'^\*\*(ACCIÓN DE TUTELA|DERECHO DE PETICIÓN)\*\*$'
    patron_titulo_seccion = r'^\*\*([IVX]+\.|[IVX]+)\s*.+\*\*$'  # I., II., III. o I, II, III
    patron_linea_negrita = r'^\*\*.+\*\*$'
    patron_guion_bajo = r'^_{3,}$'  # Líneas de firma ___________

    for i, linea in enumerate(lineas):
        linea_stripped = linea.strip()

        if not linea_stripped:
            # Línea vacía - agregar espacio pequeño
            story.append(Spacer(1, 0.15 * inch))
            continue

        # Detectar título principal (ACCIÓN DE TUTELA o DERECHO DE PETICIÓN)
        if re.match(patron_titulo_principal, linea_stripped, re.IGNORECASE):
            texto_limpio = linea_stripped.replace('**', '').strip()
            p = Paragraph(f"<b>{texto_limpio}</b>", styles['TituloPrincipal'])
            story.append(p)
            story.append(Spacer(1, 0.2 * inch))

        # Detectar títulos de sección con numeración romana (I. HECHOS, II. DERECHOS, etc.)
        elif re.match(patron_titulo_seccion, linea_stripped):
            texto_limpio = linea_stripped.replace('**', '').strip()
            p = Paragraph(f"<b>{texto_limpio}</b>", styles['TituloSeccion'])
            story.append(p)
            story.append(Spacer(1, 0.1 * inch))

        # Detectar otras líneas en negrita
        elif re.match(patron_linea_negrita, linea_stripped):
            texto_limpio = linea_stripped.replace('**', '').strip()
            # Si está en mayúsculas o es corto, podría ser subtítulo
            if linea_stripped.isupper() or len(texto_limpio) < 80:
                p = Paragraph(f"<b>{texto_limpio}</b>", styles['Encabezado'])
            else:
                p = Paragraph(f"<b>{texto_limpio}</b>", styles['Justify'])
            story.append(p)

        # Detectar líneas de firma (guiones bajos)
        elif re.match(patron_guion_bajo, linea_stripped):
            p = Paragraph(linea_stripped, styles['Encabezado'])
            story.append(p)

        # Texto normal
        else:
            # Preservar formato de listas y viñetas
            p = Paragraph(linea, styles['Justify'])
            story.append(p)

    # Construir el PDF con numeración de páginas
    doc.build(story, onFirstPage=_agregar_numero_pagina, onLaterPages=_agregar_numero_pagina)
    buffer.seek(0)

    return buffer


def generar_docx(documento_texto: str, nombre_solicitante: str = "Documento") -> BytesIO:
    """
    Genera un archivo DOCX del documento legal con formato profesional
    """
    buffer = BytesIO()

    # Crear documento
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Agregar numeración de páginas
        section.footer.is_linked_to_previous = False
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Página "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Agregar campo de número de página
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Configurar fuente del pie de página
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        footer_para.runs[0].font.name = 'Times New Roman'
        footer_para.runs[0].font.size = Pt(9)

    # Dividir el texto en líneas
    lineas = documento_texto.split('\n')

    # Patrones para detectar diferentes tipos de formato
    patron_titulo_principal = r'^\*\*(ACCIÓN DE TUTELA|DERECHO DE PETICIÓN)\*\*$'
    patron_titulo_seccion = r'^\*\*([IVX]+\.|[IVX]+)\s*.+\*\*$'  # I., II., III. o I, II, III
    patron_linea_negrita = r'^\*\*.+\*\*$'
    patron_guion_bajo = r'^_{3,}$'  # Líneas de firma ___________

    for linea in lineas:
        linea_stripped = linea.strip()

        if not linea_stripped:
            # Línea vacía
            doc.add_paragraph()
            continue

        # Detectar título principal (ACCIÓN DE TUTELA o DERECHO DE PETICIÓN)
        if re.match(patron_titulo_principal, linea_stripped, re.IGNORECASE):
            texto_limpio = linea_stripped.replace('**', '').strip()
            heading = doc.add_heading(texto_limpio, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Configurar fuente del título principal
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True

        # Detectar títulos de sección con numeración romana (I. HECHOS, II. DERECHOS, etc.)
        elif re.match(patron_titulo_seccion, linea_stripped):
            texto_limpio = linea_stripped.replace('**', '').strip()
            heading = doc.add_heading(texto_limpio, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Configurar fuente de títulos de sección
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True

        # Detectar otras líneas en negrita
        elif re.match(patron_linea_negrita, linea_stripped):
            texto_limpio = linea_stripped.replace('**', '').strip()
            p = doc.add_paragraph()

            # Si está en mayúsculas o es corto, centrarlo
            if linea_stripped.isupper() and len(texto_limpio) < 80:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(texto_limpio)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True

        # Detectar líneas de firma (guiones bajos)
        elif re.match(patron_guion_bajo, linea_stripped):
            p = doc.add_paragraph(linea_stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # Texto normal
        else:
            p = doc.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Configurar fuente
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    # Guardar en buffer
    doc.save(buffer)
    buffer.seek(0)

    return buffer
